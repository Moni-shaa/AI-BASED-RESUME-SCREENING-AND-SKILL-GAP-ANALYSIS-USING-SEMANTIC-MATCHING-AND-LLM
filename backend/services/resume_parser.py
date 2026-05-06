import PyPDF2
from docx import Document
import re
import os
from datetime import datetime
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
# Load semantic embedding model once
semantic_model = SentenceTransformer("all-MiniLM-L6-v2")

class ResumeParser:
    def __init__(self):
        # Enhanced skill patterns for better detection
        self.technical_skills = [
            # Programming Languages
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin',
            'Scala', 'R', 'MATLAB', 'Perl', 'Shell', 'Bash', 'PowerShell', 'VB.NET', 'Objective-C', 'Dart', 'Elixir',
            
            # Web Technologies
            'HTML', 'CSS', 'SCSS', 'SASS', 'React', 'Angular', 'Vue.js', 'Vue', 'Node.js', 'Express.js', 'Express',
            'Next.js', 'Nuxt.js', 'Svelte', 'jQuery', 'Bootstrap', 'Tailwind CSS', 'Material-UI', 'Chakra UI',
            
            # Databases
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server', 'Cassandra',
            'DynamoDB', 'Firebase', 'Elasticsearch', 'Neo4j', 'CouchDB', 'MariaDB',
            
            # Cloud & DevOps
            'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes', 'Jenkins', 'CI/CD', 'DevOps',
            'Terraform', 'Ansible', 'Chef', 'Puppet', 'Vagrant', 'Nginx', 'Apache', 'Linux', 'Ubuntu',
            
            # Frameworks & Libraries
            'Django', 'Flask', 'FastAPI', 'Spring', 'Spring Boot', 'Hibernate', 'Laravel', 'Symfony',
            'Rails', 'Ruby on Rails', 'ASP.NET', '.NET', 'Entity Framework', 'Xamarin',
            
            # Mobile Development
            'Android', 'iOS', 'React Native', 'Flutter', 'Ionic', 'Cordova', 'PhoneGap',
            
            # Data Science & AI
            'Machine Learning', 'Deep Learning', 'AI', 'Artificial Intelligence', 'Data Science', 'Data Analysis',
            'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn',
            'Jupyter', 'Apache Spark', 'Hadoop', 'Tableau', 'Power BI', 'Excel', 'Statistics',
            
            # Version Control & Tools
            'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN', 'Mercurial', 'JIRA', 'Confluence', 'Slack',
            'Trello', 'Asana', 'Notion', 'VS Code', 'IntelliJ', 'Eclipse', 'Visual Studio',
            
            # Methodologies
            'Agile', 'Scrum', 'Kanban', 'Waterfall', 'DevOps', 'TDD', 'BDD', 'Microservices',
            'REST API', 'GraphQL', 'SOAP', 'API Development', 'Web Services',
            
            # Soft Skills
            'Leadership', 'Team Management', 'Project Management', 'Communication', 'Problem Solving',
            'Critical Thinking', 'Analytical Skills', 'Time Management', 'Collaboration', 'Mentoring'
        ]
        
        # Education keywords
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 'college', 'institute',
            'school', 'education', 'b.s.', 'm.s.', 'b.a.', 'm.a.', 'mba', 'b.tech', 'm.tech',
            'b.e.', 'm.e.', 'diploma', 'certificate', 'certification', 'course', 'training'
        ]
    
    def extract_text(self, file_path):
        """Extract text from PDF or DOCX file"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
    def extract_text_from_file_object(self, file):
        """
        Extract text from uploaded file object (Flask/FastAPI file)
        """
        try:
           filename = file.filename.lower()

           if filename.endswith('.pdf'):
              pdf_reader = PyPDF2.PdfReader(file)
              text = ""
              for page in pdf_reader.pages:
                  extracted = page.extract_text()
                  if extracted:
                      text += extracted + "\n"
              return self._clean_text(text)

           elif filename.endswith('.docx'):
               doc = Document(file)
               text = ""
               for paragraph in doc.paragraphs:
                   text += paragraph.text + "\n"
               return self._clean_text(text)

           else:
               raise ValueError(f"Unsupported file format: {filename}")

        except Exception as e:
            raise Exception(f"Error extracting text from uploaded file: {str(e)}")  
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text while preserving structure"""
        # Preserve line breaks for better parsing
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive whitespace but keep structure
            cleaned_line = re.sub(r'\s+', ' ', line.strip())
            if cleaned_line:  # Only add non-empty lines
                cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    def extract_contact_info(self, text):
        """Advanced contact information extraction"""
        contact_info = {}
        lines = text.split('\n')
        
        # Enhanced email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        if emails:
            contact_info['email'] = emails[0]
        
        # Enhanced phone number extraction
        phone_patterns = [
            r'\+\d{1,3}[-.\s]?\d{10}',  # International format
            r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',  # (123) 456-7890
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # 123-456-7890
            r'\d{10}',  # 1234567890
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break
        
        # Advanced name extraction
        name = self._extract_name_advanced(lines)
        if name:
            contact_info['name'] = name
        
        return contact_info
    
    def _extract_name_advanced(self, lines):
        """Advanced name extraction with multiple strategies"""
        # Strategy 1: Look for name patterns in first few lines
        for i, line in enumerate(lines[:8]):
            line = line.strip()
            
            # Skip empty lines, emails, phones, addresses
            if not line or '@' in line or any(char.isdigit() for char in line if len([c for c in line if c.isdigit()]) > 3):
                continue
            
            # Skip common resume headers
            skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary', 'objective', 
                         'experience', 'education', 'skills', 'contact', 'address', 'phone', 'email']
            if any(word in line.lower() for word in skip_words):
                continue
            
            # Look for name patterns
            words = line.split()
            if 2 <= len(words) <= 4:  # Names typically have 2-4 words
                # Check if it looks like a name (mostly alphabetic)
                if all(word.replace('.', '').replace(',', '').isalpha() for word in words):
                    # Check if it's in title case or all caps
                    if line.istitle() or line.isupper():
                        return line
        
        # Strategy 2: Look for "Name:" pattern
        for line in lines[:10]:
            name_match = re.search(r'name\s*[:\-]\s*([A-Za-z\s\.]+)', line, re.IGNORECASE)
            if name_match:
                return name_match.group(1).strip()
        
        return None
    
    def extract_skills(self, text):
        """Improved skill extraction using word boundaries"""

        found_skills = []
        text_lower = text.lower()

        for skill in self.technical_skills:

             pattern = r'\b' + re.escape(skill.lower()) + r'\b'

             if re.search(pattern, text_lower):
                 found_skills.append(skill)
                 continue

             variations = self._get_skill_variations(skill)

             for variation in variations:

                 pattern = r'\b' + re.escape(variation.lower()) + r'\b'

                 if re.search(pattern, text_lower):
                     found_skills.append(skill)
                     break

        return list(set(found_skills))
    
    def _get_skill_variations(self, skill):
        """Get common variations of skill names"""
        variations = [skill]
        
        skill_variations = {
            'JavaScript': ['JS', 'Javascript', 'ECMAScript'],
            'TypeScript': ['TS', 'Typescript'],
            'Node.js': ['Node', 'NodeJS', 'Node JS'],
            'React': ['ReactJS', 'React.js'],
            'Angular': ['AngularJS', 'Angular.js'],
            'Vue.js': ['Vue', 'VueJS', 'Vue JS'],
            'C++': ['C plus plus', 'CPP'],
            'C#': ['C sharp', 'CSharp'],
            'ASP.NET': ['ASP NET', 'ASPNET'],
            'Machine Learning': ['ML', 'MachineLearning'],
            'Artificial Intelligence': ['AI'],
            'Deep Learning': ['DL', 'DeepLearning'],
            'Data Science': ['DataScience'],
            'PostgreSQL': ['Postgres', 'PostGres'],
            'MongoDB': ['Mongo'],
            'Express.js': ['Express', 'ExpressJS'],
            'REST API': ['REST', 'RESTful', 'RESTful API'],
            'GraphQL': ['Graph QL'],
        }
        
        if skill in skill_variations:
            variations.extend(skill_variations[skill])
        
        return variations
    
    
    def extract_experience_years(self, text):
        """Enhanced experience extraction"""
        # Multiple patterns for experience detection
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?(?:experience|exp)',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in',
            r'(\d+)\+?\s*years?\s*(?:working|work)',
            r'total\s*(?:of\s*)?(\d+)\+?\s*years?',
            r'over\s*(\d+)\+?\s*years?',
            r'more\s*than\s*(\d+)\+?\s*years?'
        ]
        
        text_lower = text.lower()
        max_experience = 0
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match)
                    max_experience = max(max_experience, years)
                except ValueError:
                    continue
        
        # If no explicit experience found, estimate from work history dates
        if max_experience == 0:
            max_experience = self._estimate_experience_from_dates(text)
        
        return max_experience
    
    def _estimate_experience_from_dates(self, text):
        """Estimate experience from date ranges in work history"""
        current_year = datetime.now().year
        years = re.findall(r'\b(19|20)\d{2}\b', text)
        
        if len(years) >= 2:
            years = [int(year) for year in years if 1990 <= int(year) <= current_year]
            if years:
                earliest_year = min(years)
                return min(current_year - earliest_year, 25)  # Cap at 25 years
        
        return 0
    
    def extract_education(self, text):
        """Extract education information with better parsing"""
        education_info = []
        lines = text.split('\n')
        
        # Look for education sections
        education_section_found = False
        current_education = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line indicates start of education section
            if any(keyword in line.lower() for keyword in ['education', 'academic', 'qualification']):
                education_section_found = True
                continue
            
            # If we're in education section or line contains education keywords
            if education_section_found or any(keyword in line.lower() for keyword in self.education_keywords):
                # Skip section headers
                if line.lower() in ['education', 'academic background', 'qualifications']:
                    continue
                
                # Look for degree patterns
                degree_patterns = [
                    r'(bachelor|master|phd|doctorate|diploma|certificate).*?(?:in|of)\s*([^,\n]+)',
                    r'(b\.?[aes]\.?|m\.?[aes]\.?|phd|mba|m\.?tech|b\.?tech).*?(?:in|of)?\s*([^,\n]+)',
                    r'(degree)\s*(?:in|of)\s*([^,\n]+)'
                ]
                
                for pattern in degree_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        degree_type = match.group(1)
                        field = match.group(2) if len(match.groups()) > 1 else ""
                        education_info.append(f"{degree_type} {field}".strip())
                        break
                else:
                    # If no pattern matched but contains education keywords, add the line
                    if any(keyword in line.lower() for keyword in self.education_keywords):
                        education_info.append(line)
        
        return education_info if education_info else ["Education information not clearly specified"]
    # -----------------------------------
    # PRIORITY + WEIGHTED MATCH LOGIC
    # -----------------------------------

import re

def detect_priority(skill, jd_text):
    jd = jd_text.lower()
    skill = skill.lower()

    high_keywords = ["must", "required", "mandatory", "strong knowledge"]
    low_keywords = ["preferred", "nice to have", "plus"]

    # find occurrences of the skill in JD
    matches = [m.start() for m in re.finditer(skill, jd)]

    for pos in matches:
        # take a window around the skill (40 chars before/after)
        window = jd[max(0, pos-40): pos+40]

        for word in high_keywords:
            if word in window:
                return "high"

        for word in low_keywords:
            if word in window:
                return "low"

    # default
    return "medium"


def get_weight(priority):
    weights = {
        "high": 5,
        "medium": 3,
        "low": 1
    }
    return weights.get(priority, 3)


def calculate_weighted_match(jd_skills, resume_skills, jd_text,resume_text):
    total_weight = 0
    matched_weight = 0
    matched_skills = []
    missing_skills = []
    missing_critical = []

    priority_counts = {
       "high": 0,
       "medium": 0,
       "low": 0,
       "missing": 0
   }

    for skill in jd_skills:

        priority = detect_priority(skill, jd_text)
        weight = get_weight(priority)

        total_weight += weight

        if skill in resume_skills:

           matched_weight += weight
           matched_skills.append(skill)

           priority_counts[priority] += 1

        else:

           missing_skills.append(skill)
           priority_counts["missing"] += 1

           if priority == "high":
              missing_critical.append(skill)

    weighted_score = (matched_weight / total_weight) * 100 if total_weight else 0

    semantic_score = calculate_semantic_similarity(resume_text, jd_text)

    final_score = round(
        (0.6 * semantic_score) + (0.4 * weighted_score),
        2
    )

    

    return {
    "match_score": float(round(weighted_score,2)),
    "semantic_score": float(semantic_score),
    "final_score": float(final_score),
    "matched_skills": matched_skills,
    "missing_skills": missing_skills,
    "missing_critical": missing_critical,
    "priority_distribution": priority_counts
}


def decide_shortlist(final_score, missing_critical):
    if final_score >= 70:
        return 1, "Strong match with Job Description"

    if missing_critical:
        return 0, f"Missing critical skills: {', '.join(missing_critical)}"

    if final_score >= 50:
        return 1, "Moderate match — shortlisted for review"

    return 0, "Low match score"

def calculate_semantic_similarity(resume_text, jd_text):
    """
    Compute semantic similarity between resume and job description
    using sentence embeddings.
    """

    if not resume_text or not jd_text:
        return 0

    embeddings = semantic_model.encode([resume_text, jd_text])

    similarity = cosine_similarity(
        [embeddings[0]],
        [embeddings[1]]
    )[0][0]

    return round(float(similarity * 100), 2)
    

